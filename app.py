from flask import Flask, render_template, request, redirect, url_for, send_file
import os
from werkzeug.utils import secure_filename
from parsers.extract_data import extract_resume_data
from docx import Document  # Required for resume continuation

app = Flask(__name__)

# Config upload and output directories
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Ensure folders exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Route: Home page (index.html)
@app.route('/')
def index():
    return render_template('index.html')  # <- now points to index.html

# Route: Upload page (GET and POST)
@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if request.method == 'GET':
        return render_template('upload.html')

    # Handle POST (file upload)
    if 'resume' not in request.files:
        return "No file part", 400

    file = request.files['resume']
    if file.filename == '':
        return "No selected file", 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    # Extract data from resume
    extracted_data = extract_resume_data(filepath)

    return render_template('form.html', data=extracted_data)

# Route: Resume template preview
@app.route('/preview')
def preview():
    name = request.args.get('name', '')
    email = request.args.get('email', '')
    linkedin = request.args.get('linkedin', '')
    github = request.args.get('github', '')
    summary = request.args.get('summary', '')
    template = request.args.get('template', 'template1')

    return render_template('preview.html', name=name, email=email, linkedin=linkedin,
                           github=github, summary=summary, template=template)

# Route: Template selection
@app.route('/template')
def template_selection():
    name = request.args.get('name', '')
    email = request.args.get('email', '')
    linkedin = request.args.get('linkedin', '')
    github = request.args.get('github', '')
    summary = request.args.get('summary', '')

    return render_template('template.html', name=name, email=email, linkedin=linkedin,
                           github=github, summary=summary)
@app.route('/builder')
def builder():
    return render_template('builder.html')


# Route: Continue editing existing resume
@app.route('/continue-existing', methods=['POST'])
def continue_existing():
    filename = request.form['filename']
    filepath = os.path.join(UPLOAD_FOLDER, filename)

    doc = Document(filepath)
    existing_text = "\n".join(p.text.lower() for p in doc.paragraphs)

    def add_if_not_present(label, value):
        if value and value.lower() not in existing_text:
            doc.add_paragraph(f"{label}: {value}")

    add_if_not_present("Name", request.form['name'])
    add_if_not_present("Email", request.form['email'])
    add_if_not_present("LinkedIn", request.form['linkedin'])
    add_if_not_present("GitHub", request.form['github'])
    add_if_not_present("Summary", request.form['summary'])

    output_path = os.path.join(OUTPUT_FOLDER, f'modified_{filename}')
    doc.save(output_path)

    return send_file(output_path, as_attachment=True)

# Start the app
if __name__ == '__main__':
    app.run(debug=True)
